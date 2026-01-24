"""
Document Reader Module
Extracts text from PDF and DOCX files using pdfminer, python-docx, and OCR fallback.
"""

import io
import os
from typing import Optional

try:
    from pdfminer.high_level import extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def setup_tesseract():
    """Configure Tesseract OCR path based on OS."""
    if os.name == 'nt':  # Windows
        tesseract_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(os.getlogin()),
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return True
    return False


# Try to setup Tesseract on module load
if OCR_AVAILABLE:
    setup_tesseract()


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes.
    
    Args:
        docx_bytes: Raw bytes of the DOCX file
    
    Returns:
        Extracted text string
    """
    if not DOCX_AVAILABLE:
        return ""
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        return '\n'.join(paragraphs)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF bytes using multiple methods.
    
    1. Try normal text extraction (digital PDFs)
    2. If empty -> fallback to OCR (scanned PDFs)
    
    Args:
        pdf_bytes: Raw bytes of the PDF file
    
    Returns:
        Extracted text string
    """
    if not pdf_bytes:
        return ""

    # Try pdfminer first (for digital/text-based PDFs)
    if PDFMINER_AVAILABLE:
        try:
            text = extract_text(io.BytesIO(pdf_bytes))
            if text and text.strip() and len(text.strip()) > 50:
                return clean_extracted_text(text)
        except Exception as e:
            print(f"pdfminer extraction failed: {e}")

    # OCR fallback (for scanned PDFs)
    if OCR_AVAILABLE:
        try:
            images = convert_from_bytes(pdf_bytes, dpi=200)
            ocr_text = ""
            for i, img in enumerate(images):
                page_text = pytesseract.image_to_string(img, lang='eng')
                ocr_text += f"\n--- Page {i+1} ---\n{page_text}"
            
            if ocr_text.strip():
                return clean_extracted_text(ocr_text)
        except Exception as e:
            print(f"OCR extraction failed: {e}")

    return ""


def extract_text_from_pdf_file(file_path: str) -> str:
    """
    Extract text from a PDF file path.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        Extracted text string
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    with open(file_path, 'rb') as f:
        pdf_bytes = f.read()
    
    return extract_text_from_pdf(pdf_bytes)


def clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text by removing excessive whitespace.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple newlines with double newline
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    
    return '\n'.join(lines).strip()


def get_pdf_info(pdf_bytes: bytes) -> dict:
    """
    Get information about a PDF file.
    
    Args:
        pdf_bytes: Raw bytes of the PDF file
    
    Returns:
        Dictionary with PDF metadata
    """
    info = {
        "size_bytes": len(pdf_bytes),
        "size_kb": round(len(pdf_bytes) / 1024, 2),
        "pdfminer_available": PDFMINER_AVAILABLE,
        "ocr_available": OCR_AVAILABLE,
    }
    
    # Try to get page count
    if PDFMINER_AVAILABLE:
        try:
            from pdfminer.pdfpage import PDFPage
            pages = list(PDFPage.get_pages(io.BytesIO(pdf_bytes)))
            info["page_count"] = len(pages)
        except Exception:
            info["page_count"] = "Unknown"
    
    return info

